# -*- coding: utf-8 -*-
"""Generates docs/MLB_Model_V2_Stats_Catalog.docx — a read-through catalog of
every stat the pipeline can pull, grouped by source. Run after any change to
app/stat_definitions.py to keep the doc in sync:

    venv\\Scripts\\python.exe app\\generate_stats_catalog.py
"""
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from app.stat_definitions import SECTIONS  # noqa: E402

OUT_PATH = Path(__file__).resolve().parent.parent / "docs" / "MLB_Model_V2_Stats_Catalog.docx"

HEADER_FILL = "2F5496"


def _shade_cell(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def build():
    doc = Document()

    title = doc.add_heading("MLB Model V2 — Stat Catalog", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Every stat this pipeline can pull, as of {date.today():%B %d, %Y}")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "This catalog lists every stat currently collected by the R/ data pipeline and every "
        "feature engineered from them in features/builder.py. Each section below covers one raw "
        "data source (or one engineered-feature group) and lists the stats it provides, what they "
        "mean, and their scale. Team and starting-pitcher season stats are collected back to 2020, "
        "batter/fielding/lineup data back to 2021 — the search app built on top of this data is "
        "currently scoped to the 2026 season only."
    )

    doc.add_page_break()

    # Table of contents (simple manual list, not a live TOC field)
    doc.add_heading("Contents", level=1)
    for s in SECTIONS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s["title"])

    doc.add_page_break()

    for s in SECTIONS:
        doc.add_heading(s["title"], level=1)

        meta = doc.add_paragraph()
        meta.add_run("Source: ").bold = True
        meta.add_run(s["source"])
        meta2 = doc.add_paragraph()
        meta2.add_run("Grain: ").bold = True
        meta2.add_run(s["grain"])

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["Column", "Stat", "Description", "Units / Scale"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(hdr[i], HEADER_FILL)

        for key, label, desc, units in s["stats"]:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = label
            row[2].text = desc
            row[3].text = units

        doc.add_paragraph()

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
